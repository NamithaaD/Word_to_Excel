import streamlit as st
import pandas as pd
import re
from datetime import datetime
from docx import Document
from io import BytesIO
from PIL import Image

# === PAGE CONFIG ===
st.set_page_config(page_title="Care Center", layout="centered")

# === LOGO + COMPANY NAME SIDE BY SIDE ===
logo = Image.open("nupie.png")
col1, col2 = st.columns([1, 6])
with col1:
    st.image(logo, width=80)
with col2:
    st.markdown(
        "<h2 style='margin-bottom: 0;'>CARE CENTER</h2>",
        unsafe_allow_html=True
    )

# === TITLE & DESCRIPTION ===
st.title("AMC Word Extractor")
st.markdown(
    "Upload one or more `.docx` files to extract AMC contract details and download as Excel. "
    "_(No data is stored)_"
)

# === FUNCTION TO EXTRACT TEXT AND TABLES ===
def extract_text_from_docx(file):
    try:
        doc = Document(file)
        text = '\n'.join(p.text for p in doc.paragraphs)
        table_data = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            table_data.append(rows)
        return text, table_data
    except Exception as e:
        st.warning(f"Error reading document: {e}")
        return '', []

# === CONTRACT NUMBER FROM FILENAME ===
def extract_contract_number_from_filename(filename):
    m = re.search(r'CC(\d+)_', filename)
    return m.group(1) if m else ''

# === GET FIRST VALID DATE ===
def extract_first_valid_date(text):
    for date_str in re.findall(r'\b(\d{2}\.\d{2}\.\d{4})\b', text):
        try:
            return datetime.strptime(date_str, '%d.%m.%Y').strftime('%d-%m-%Y')
        except:
            continue
    return ''

# === MAIN EXTRACTION FUNCTION ===
def extract_details(text, contract_no, table_data):
    fields = {
        "Contract Type": "",
        "Contract No": contract_no,
        "Customer Name": "",
        "Contact Number": "",
        "Address": "",
        "Location": "",
        "Unit Details": "",
        "Amount": "",
        "Contract Date": "",
        "Amc Start Date": "",
        "Amc End Date": "",
        "Service Frequency": 4,
        "1st Service Month": "",
        "2nd Service Month": "",
        "3rd Service Month": "",
        "4th Service Month": ""
    }

    def find(pat, fb=""):
        m = re.search(pat, text, re.IGNORECASE | re.DOTALL)
        return m.group(1).strip() if m else fb

    if re.search(r'ANNUAL MAINTENANCE CONTRACT', text, re.IGNORECASE):
        fields["Contract Type"] = "ANNUAL MAINTENANCE CONTRACT"
    elif re.search(r'LABOUR MAINTENANCE CONTRACT', text, re.IGNORECASE):
        fields["Contract Type"] = "LABOUR MAINTENANCE CONTRACT"

    fields["Customer Name"] = find(r'Customer Name\s*:\s*(.*?)\n')
    fields["Contact Number"] = find(r'Contact Number\s*:\s*(.*?)\n').split("UNIT DETAILS")[0].strip()

    addr = re.search(r'Address\s*:\s*(.*?)(?:Contact Number|Location|Unit Details)', text, re.IGNORECASE | re.DOTALL)
    if addr:
        fields["Address"] = re.sub(r'\s+', ' ', addr.group(1)).strip()

    loc = re.search(r'Location\s*:\s*(.*?)\n', text, re.IGNORECASE)
    if loc and 'CONTACT NUMBER' not in loc.group(1).upper():
        fields["Location"] = loc.group(1).strip()

    fields["Contract Date"] = extract_first_valid_date(text)

    amount_match = re.search(r'AMOUNT\s*[:\-]?\s*(?:Rs\.?|₹)?\s*([\d,]+(?:\.\d{1,2})?)', text, re.IGNORECASE)
    if not amount_match:
        amount_match = re.search(r'(?:Rs\.?|₹)?\s*([\d,]+(?:\.\d{1,2})?)\s*/-', text)
    fields["Amount"] = amount_match.group(1).strip() if amount_match else ""

    amc = re.search(r'AMC PERIOD\s*[:\-]?\s*(\d{2}\.\d{2}\.\d{4})\s*(?:to|-)\s*(\d{2}\.\d{2}\.\d{4})', text, re.IGNORECASE)
    if amc:
        try:
            start = datetime.strptime(amc.group(1), "%d.%m.%Y")
            end = datetime.strptime(amc.group(2), "%d.%m.%Y")
            fields["Amc Start Date"] = start.strftime('%d-%m-%Y')
            fields["Amc End Date"] = end.strftime('%d-%m-%Y')
            for i in range(4):
                svc = start + pd.DateOffset(months=2 + i * 3)
                label = f"{i+1}st Service Month" if i == 0 else f"{i+1}nd Service Month" if i == 1 else f"{i+1}rd Service Month" if i == 2 else f"{i+1}th Service Month"
                fields[label] = svc.strftime('%B %Y')
        except:
            pass

    unit_details = ""
    extracted = False
    for table in table_data:
        if len(table) > 1:
            headers = [c.lower() for c in table[0]]
            if any(h in headers for h in ('brand','range','rate','qty','amount','ton')):
                for row in table[1:]:
                    row_text = " | ".join(cell for cell in row if cell.strip())
                    if row_text and not re.search(r'\bTOTAL\b', row_text, re.IGNORECASE):
                        unit_details += row_text + "\n"
                extracted = True
                break

    if not extracted:
        match = re.search(r'UNIT DETAILS\s*:?(.+?)(?=AMOUNT|AMC PERIOD|CONTRACT NO|DATE|SERVICE|TERMS & CONDITIONS|$)', text, re.IGNORECASE | re.DOTALL)
        if match:
            for line in match.group(1).splitlines():
                clean_line = line.strip()
                if clean_line and not re.search(r'Terms & Conditions|routine service', clean_line, re.IGNORECASE):
                    unit_details += clean_line + "\n"

    fields["Unit Details"] = unit_details.strip()
    return fields

# === FILE UPLOADER ===
uploaded_files = st.file_uploader("Upload DOCX files", type="docx", accept_multiple_files=True)

# === PROCESS FILES ===
if uploaded_files:
    extracted_data = []
    for file in uploaded_files:
        txt, tables = extract_text_from_docx(file)
        if txt.strip():
            contract_no = extract_contract_number_from_filename(file.name)
            details = extract_details(txt, contract_no, tables)
            extracted_data.append(details)

    if extracted_data:
        df = pd.DataFrame(extracted_data)

        st.subheader(" Preview Extracted Data")
        st.dataframe(df, use_container_width=True)

        output = BytesIO()
        df.to_excel(output, index=False, engine='openpyxl')
        output.seek(0)

        st.success(" Extraction complete!")
        st.download_button(
            label=" Download Excel File",
            data=output,
            file_name=f"amc_output_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

# === FOOTER ===
st.markdown(
    """<hr style="margin-top: 2em;">
    <div style='text-align: center; font-size: 0.9em; color: gray;'>
        Built by <a href="https://nu-pie.ai/" target="_blank" style="text-decoration: none; color: #636363;">
        Nu-pie Analytics</a>
    </div>
    """, unsafe_allow_html=True
)
